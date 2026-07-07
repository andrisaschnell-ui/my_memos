from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse
import pypdf
import fitz # PyMuPDF
from docx import Document
import io
import base64
import os
import httpx
import logging

logger = logging.getLogger("document")
router = APIRouter(prefix="/document", tags=["document"])

# Load NVIDIA API Key with local file fallback
NVIDIA_API_KEY = os.getenv("NVIDIA-KEY") or os.getenv("NVIDIA_API_KEY")
if not NVIDIA_API_KEY:
    try:
        secrets_path = r"c:\Users\Andrisa\Documents\Projects\mem_assist\secrets\nvapi-pgmqbRBYjl1Htlp0uOz0hz_NE3fp9.txt"
        if os.path.exists(secrets_path):
            with open(secrets_path, "r") as f:
                NVIDIA_API_KEY = f.read().strip()
    except Exception as e:
        logger.warning(f"Failed to load local NVIDIA key in document router: {e}")

NVIDIA_URL = "https://integrate.api.nvidia.com/v1/chat/completions"

async def ocr_page_with_nvidia(img_bytes: bytes, page_num: int) -> str:
    if not NVIDIA_API_KEY:
        raise Exception("NVIDIA API key not configured")
        
    b64_image = base64.b64encode(img_bytes).decode("utf-8")
    
    headers = {
        "Authorization": f"Bearer {NVIDIA_API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "meta/llama-3.2-11b-vision-instruct",
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": "Analyze this page image from a PDF document and perform OCR to extract all readable text word-for-word. Do your best to preserve paragraphs, lists, sections, bullet points, headers, and basic table alignments. Output only the exact extracted text from the page. Do not include any greeting, explanation, or meta notes."
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{b64_image}"
                        }
                    }
                ]
            }
        ],
        "max_tokens": 2048,
        "temperature": 0.1,
        "top_p": 0.7
    }
    
    async with httpx.AsyncClient() as client:
        resp = await client.post(NVIDIA_URL, json=payload, headers=headers, timeout=60.0)
        if resp.status_code != 200:
            raise Exception(f"NVIDIA API error on page {page_num}: {resp.status_code} - {resp.text}")
        
        res_json = resp.json()
        return res_json["choices"][0]["message"]["content"]

@router.post("/pdf-to-docx")
async def pdf_to_docx(
    file: UploadFile = File(...),
    mode: str = Form("fast") # "fast" or "ocr"
):
    try:
        file_bytes = await file.read()
        if not file_bytes:
            raise HTTPException(status_code=400, detail="File is empty")
            
        extracted_text = ""
        
        if mode == "ocr":
            # Image OCR Mode using PyMuPDF and Llama Vision Model
            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Failed to open PDF file: {e}")
                
            num_pages = len(doc)
            if num_pages == 0:
                raise HTTPException(status_code=400, detail="PDF has 0 pages")
                
            for idx in range(num_pages):
                page = doc.load_page(idx)
                # 150 DPI is a sweet spot for Llama vision resolution and performance
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")
                
                try:
                    page_text = await ocr_page_with_nvidia(img_bytes, idx + 1)
                except Exception as e:
                    page_text = f"[OCR Failed on Page {idx + 1}: {str(e)}]"
                    
                extracted_text += f"\n--- Page {idx + 1} ---\n{page_text}\n"
                
            doc.close()
        else:
            # Fast Text-selectable Parsing Mode using PyPDF
            try:
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                for idx, page in enumerate(reader.pages):
                    page_text = page.extract_text() or ""
                    extracted_text += f"\n--- Page {idx + 1} ---\n{page_text}\n"
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Failed to parse text from PDF: {e}")

        # Assemble DOCX Document using python-docx
        docx_doc = Document()
        # Add basic document header
        h = docx_doc.add_heading(level=1)
        run = h.add_run("Extracted Document Text")
        run.font.name = 'Arial'
        
        # Split text into paragraphs and write into DOCX
        paragraphs = extracted_text.split("\n")
        for p in paragraphs:
            p_clean = p.strip()
            if p_clean:
                if p_clean.startswith("--- Page") and p_clean.endswith("---"):
                    docx_doc.add_heading(p_clean, level=2)
                else:
                    docx_doc.add_paragraph(p_clean)
                    
        # Save to memory stream
        docx_stream = io.BytesIO()
        docx_doc.save(docx_stream)
        docx_stream.seek(0)
        
        docx_base64 = base64.b64encode(docx_stream.read()).decode("utf-8")
        
        return JSONResponse(content={
            "success": True,
            "text": extracted_text,
            "docx_base64": docx_base64,
            "filename": f"{os.path.splitext(file.filename)[0]}.docx"
        })
        
    except Exception as err:
        logger.error(f"Error converting PDF to DOCX: {err}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"PDF extraction error: {str(err)}")
