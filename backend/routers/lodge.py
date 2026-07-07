from fastapi import APIRouter, Depends, HTTPException, status, Header
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update, delete, func, desc
from datetime import date, datetime, timedelta
from typing import List, Optional
import uuid

from database import get_db
from models import Guest, Reservation, LodgeTask, Incident, DailyLog, Recording
from schemas import (
    GuestCreate, GuestOut,
    ReservationCreate, ReservationOut,
    LodgeTaskCreate, LodgeTaskOut,
    IncidentCreate, IncidentOut,
    DailyLogCreate, DailyLogOut
)

router = APIRouter(prefix="/lodge", tags=["lodge"])

async def get_active_lodge_id(x_lodge_id: Optional[str] = Header(None, alias="X-Lodge-Id")) -> Optional[uuid.UUID]:
    if not x_lodge_id:
        return None
    try:
        return uuid.UUID(x_lodge_id)
    except ValueError:
        return None

async def get_current_user_email(x_user_email: Optional[str] = Header(None, alias="X-User-Email")) -> Optional[str]:
    """Extract and normalize the caller's email from the X-User-Email header."""
    if not x_user_email:
        return None
    return x_user_email.strip().lower()

def parse_date(date_str: Optional[str]) -> Optional[date]:
    if not date_str:
        return None
    normalized = str(date_str).replace("/", "-").strip()
    if not normalized or normalized.lower() in ("null", "yyyy-mm-dd", "yyyy/mm/dd"):
        return None
    from datetime import datetime
    for fmt in ("%Y-%m-%d", "%Y-%m-%d %H:%M:%S", "%Y/%m/%d", "%d-%m-%Y", "%d/%m/%Y"):
        try:
            return datetime.strptime(normalized, fmt).date()
        except ValueError:
            continue
    try:
        parts = normalized.split("-")
        if len(parts) == 3:
            y = int(parts[0])
            m = int(parts[1])
            d = int(parts[2][:2])
            from datetime import date as date_type
            return date_type(y, m, d)
    except Exception:
        pass
    return None

# -----------------------------------------------------------------------------
# Module 8 Backend: Dashboard Summary Endpoint
# -----------------------------------------------------------------------------
@router.get("/dashboard")
async def get_dashboard_summary(
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    today = date.today()

    # Check-ins today
    checkins_stmt = select(Reservation).join(Guest).where(Reservation.check_in == today)
    if lodge_id:
        checkins_stmt = checkins_stmt.where(Guest.lodge_id == lodge_id)
    checkins_res = await db.execute(checkins_stmt)
    checkins = checkins_res.scalars().all()

    # Check-outs today
    checkouts_stmt = select(Reservation).join(Guest).where(Reservation.check_out == today)
    if lodge_id:
        checkouts_stmt = checkouts_stmt.where(Guest.lodge_id == lodge_id)
    checkouts_res = await db.execute(checkouts_stmt)
    checkouts = checkouts_res.scalars().all()

    # In-house tonight (check_in <= today < check_out and status != cancelled)
    inhouse_stmt = select(Reservation).join(Guest).where(
        Reservation.check_in <= today,
        Reservation.check_out > today,
        Reservation.status != "cancelled"
    )
    if lodge_id:
        inhouse_stmt = inhouse_stmt.where(Guest.lodge_id == lodge_id)
    inhouse_res = await db.execute(inhouse_stmt)
    inhouse = inhouse_res.scalars().all()
    
    inhouse_guest_ids = [r.guest_id for r in inhouse if r.guest_id]
    guests_map = {}
    if inhouse_guest_ids:
        g_res = await db.execute(select(Guest).where(Guest.id.in_(inhouse_guest_ids)))
        for g in g_res.scalars().all():
            guests_map[g.id] = g.full_name

    inhouse_list = []
    total_inhouse_count = 0
    for r in inhouse:
        gname = guests_map.get(r.guest_id, "Unknown Guest")
        inhouse_list.append({"guest_name": gname, "room": r.room_or_unit, "adults": r.num_adults, "children": r.num_children})
        total_inhouse_count += (r.num_adults or 1) + (r.num_children or 0)

    # Tasks due today or overdue
    tasks_stmt = select(LodgeTask).where(
        LodgeTask.is_complete == False,
        LodgeTask.due_date <= today
    )
    if lodge_id:
        tasks_stmt = tasks_stmt.where(LodgeTask.lodge_id == lodge_id)
    tasks_res = await db.execute(tasks_stmt.order_by(LodgeTask.due_date.asc()))
    tasks_due = tasks_res.scalars().all()

    # Open Incidents
    incidents_stmt = select(Incident).where(Incident.is_resolved == False)
    if lodge_id:
        incidents_stmt = incidents_stmt.where(Incident.lodge_id == lodge_id)
    incidents_res = await db.execute(incidents_stmt.order_by(
        desc(Incident.severity == "critical"),
        desc(Incident.created_at)
    ))
    open_incidents = incidents_res.scalars().all()

    # Recent Memos (from recordings table where type='memo')
    memos_stmt = select(Recording).where(Recording.type == "memo")
    if lodge_id:
        memos_stmt = memos_stmt.where(Recording.lodge_id == lodge_id)
    memos_res = await db.execute(memos_stmt.order_by(desc(Recording.created_at)).limit(5))
    recent_memos = memos_res.scalars().all()

    return {
        "today": today.isoformat(),
        "checkins": [{"room": r.room_or_unit, "guest_id": r.guest_id} for r in checkins],
        "checkouts": [{"room": r.room_or_unit, "guest_id": r.guest_id} for r in checkouts],
        "inhouse_count": total_inhouse_count,
        "inhouse_list": inhouse_list,
        "tasks_due_count": len(tasks_due),
        "tasks_due": [
            {"id": str(t.id), "title": t.title, "area": t.area, "assigned_to": t.assigned_to, "due_date": t.due_date.isoformat() if t.due_date else None}
            for t in tasks_due[:8]
        ],
        "incidents_count": len(open_incidents),
        "open_incidents": [
            {"id": str(i.id), "title": i.title, "area": i.area, "severity": i.severity}
            for i in open_incidents[:8]
        ],
        "recent_memos": [
            {"id": str(m.id), "summary": m.summary, "transcript": m.transcript, "created_at": m.created_at.isoformat() if m.created_at else None}
            for m in recent_memos
        ]
    }


# -----------------------------------------------------------------------------
# Guests CRUD
# -----------------------------------------------------------------------------
@router.get("/guests", response_model=List[GuestOut])
async def list_guests(
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id),
    user_email: Optional[str] = Depends(get_current_user_email)
):
    stmt = select(Guest)
    if user_email:
        stmt = stmt.where(Guest.user_email == user_email)
    if lodge_id:
        stmt = stmt.where(Guest.lodge_id == lodge_id)
    res = await db.execute(stmt.order_by(Guest.full_name.asc()))
    guests = res.scalars().all()
    for g in guests:
        g.has_passport_image = bool(g.passport_image)
    return guests

@router.post("/guests", response_model=GuestOut, status_code=201)
async def create_guest(
    data: GuestCreate,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id),
    user_email: Optional[str] = Depends(get_current_user_email)
):
    check_in = parse_date(data.check_in)
    check_out = parse_date(data.check_out)
    room_or_unit = data.room_or_unit or "Chalet"
    guest_data = data.model_dump(exclude={"check_in", "check_out", "room_or_unit"})
    
    # Safely parse date fields
    guest_data["date_of_birth"] = parse_date(guest_data.get("date_of_birth"))
    guest_data["date_of_issue"] = parse_date(guest_data.get("date_of_issue"))
    guest_data["date_of_expiry"] = parse_date(guest_data.get("date_of_expiry"))
    guest_data["visa_validity"] = parse_date(guest_data.get("visa_validity"))
    
    # Always stamp the owner email from the authenticated header
    if user_email:
        guest_data["user_email"] = user_email
    
    g = Guest(**guest_data)
    if lodge_id:
        g.lodge_id = lodge_id
        
    db.add(g)
    await db.commit()
    await db.refresh(g)
    
    if check_in and check_out:
        from models import Reservation
        res = Reservation(
            guest_id=g.id,
            check_in=check_in,
            check_out=check_out,
            room_or_unit=room_or_unit,
            status="confirmed",
            source="direct"
        )
        db.add(res)
        await db.commit()
        
    g.has_passport_image = bool(g.passport_image)
    return g

@router.put("/guests/{guest_id}", response_model=GuestOut)
async def update_guest(
    guest_id: uuid.UUID,
    data: GuestCreate,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id),
    user_email: Optional[str] = Depends(get_current_user_email)
):
    stmt = select(Guest).where(Guest.id == guest_id)
    if user_email:
        stmt = stmt.where(Guest.user_email == user_email)
    if lodge_id:
        stmt = stmt.where(Guest.lodge_id == lodge_id)
    res = await db.execute(stmt)
    g = res.scalar_one_or_none()
    if not g:
        raise HTTPException(status_code=404, detail="Guest not found")
        
    check_in = parse_date(data.check_in)
    check_out = parse_date(data.check_out)
    room_or_unit = data.room_or_unit or "Chalet"
    guest_data = data.model_dump(exclude={"check_in", "check_out", "room_or_unit"})
    
    # Safely parse date fields
    guest_data["date_of_birth"] = parse_date(guest_data.get("date_of_birth"))
    guest_data["date_of_issue"] = parse_date(guest_data.get("date_of_issue"))
    guest_data["date_of_expiry"] = parse_date(guest_data.get("date_of_expiry"))
    guest_data["visa_validity"] = parse_date(guest_data.get("visa_validity"))
    
    # Preserve owner email
    if user_email:
        guest_data["user_email"] = user_email
    
    for k, v in guest_data.items():
        setattr(g, k, v)
    await db.commit()
    await db.refresh(g)
    
    if check_in and check_out:
        from models import Reservation
        res_stmt = await db.execute(select(Reservation).where(Reservation.guest_id == g.id).order_by(Reservation.created_at.desc()))
        existing_res = res_stmt.scalars().first()
        if existing_res:
            existing_res.check_in = check_in
            existing_res.check_out = check_out
            existing_res.room_or_unit = room_or_unit
        else:
            new_res = Reservation(
                guest_id=g.id,
                check_in=check_in,
                check_out=check_out,
                room_or_unit=room_or_unit,
                status="confirmed",
                source="direct"
            )
            db.add(new_res)
        await db.commit()
        
    g.has_passport_image = bool(g.passport_image)
    return g

@router.delete("/guests/{guest_id}", status_code=204)
async def delete_guest(
    guest_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id),
    user_email: Optional[str] = Depends(get_current_user_email)
):
    stmt = select(Guest).where(Guest.id == guest_id)
    if user_email:
        stmt = stmt.where(Guest.user_email == user_email)
    if lodge_id:
        stmt = stmt.where(Guest.lodge_id == lodge_id)
    res = await db.execute(stmt)
    g = res.scalar_one_or_none()
    if g:
        await db.delete(g)
        await db.commit()
    return None


# -----------------------------------------------------------------------------
# Reservations CRUD
# -----------------------------------------------------------------------------
@router.get("/reservations")
async def list_reservations(
    upcoming: Optional[int] = None,
    status_filter: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id),
    user_email: Optional[str] = Depends(get_current_user_email)
):
    query = select(Reservation).join(Guest)
    if user_email:
        query = query.where(Guest.user_email == user_email)
    if lodge_id:
        query = query.where(Guest.lodge_id == lodge_id)
        
    query = query.order_by(Reservation.check_in.asc())
    if upcoming:
        today = date.today()
        end_date = today + timedelta(days=upcoming)
        query = query.where(Reservation.check_in >= today, Reservation.check_in <= end_date)
    if status_filter:
        query = query.where(Reservation.status == status_filter)
    
    res = await db.execute(query)
    reservations = res.scalars().all()
    
    out = []
    for r in reservations:
        g_data = None
        if r.guest_id:
            gres = await db.execute(select(Guest).where(Guest.id == r.guest_id))
            gobj = gres.scalar_one_or_none()
            if gobj:
                g_data = {
                    "id": str(gobj.id),
                    "full_name": gobj.full_name,
                    "phone": gobj.phone,
                    "email": gobj.email
                }
        out.append({
            "id": str(r.id),
            "guest_id": str(r.guest_id) if r.guest_id else None,
            "guest": g_data,
            "room_or_unit": r.room_or_unit,
            "check_in": r.check_in.isoformat() if r.check_in else None,
            "check_out": r.check_out.isoformat() if r.check_out else None,
            "num_adults": r.num_adults,
            "num_children": r.num_children,
            "rate_per_night_usd": float(r.rate_per_night_usd or 0),
            "total_usd": float(r.total_usd or 0),
            "deposit_paid": r.deposit_paid,
            "status": r.status,
            "source": r.source,
            "notes": r.notes
        })
    return out

@router.get("/immigration-report")
async def get_immigration_report(
    start_date: date,
    end_date: date,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id),
    user_email: Optional[str] = Depends(get_current_user_email)
):
    query = select(Reservation).join(Guest)
    if user_email:
        query = query.where(Guest.user_email == user_email)
    if lodge_id:
        query = query.where(Guest.lodge_id == lodge_id)
        
    query = query.where(Reservation.check_in >= start_date, Reservation.check_in <= end_date)
    query = query.order_by(Reservation.check_in.asc())
    
    res = await db.execute(query)
    reservations = res.scalars().all()
    
    out = []
    for idx, r in enumerate(reservations):
        gobj = None
        if r.guest_id:
            gres = await db.execute(select(Guest).where(Guest.id == r.guest_id))
            gobj = gres.scalar_one_or_none()
            
        if gobj:
            out.append({
                "no": idx + 1,
                "guest_id": str(gobj.id),
                "reservation_id": str(r.id),
                "full_name": gobj.full_name,
                "nationality": gobj.nationality,
                "passport_number": gobj.passport_number,
                "date_of_issue": gobj.date_of_issue.isoformat() if gobj.date_of_issue else None,
                "arrived_from": gobj.arrived_from,
                "visa_number": gobj.visa_number,
                "visa_validity": gobj.visa_validity.isoformat() if gobj.visa_validity else None,
                "check_in": r.check_in.isoformat() if r.check_in else None,
                "check_out": r.check_out.isoformat() if r.check_out else None,
            })
    return out

@router.get("/immigration-report/docx")
async def get_immigration_report_docx(
    start_date: date,
    end_date: date,
    lodge_name: Optional[str] = None,
    lodge_location: Optional[str] = None,
    language: Optional[str] = "PT",
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id),
    user_email: Optional[str] = Depends(get_current_user_email)
):
    """Generate and stream a DOCX Boletim de Alojamento / Immigration List."""
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from fastapi.responses import StreamingResponse
    import io

    # Fetch reservations
    query = select(Reservation).join(Guest)
    if user_email:
        query = query.where(Guest.user_email == user_email)
    if lodge_id:
        query = query.where(Guest.lodge_id == lodge_id)
    query = query.where(Reservation.check_in >= start_date, Reservation.check_in <= end_date)
    query = query.order_by(Reservation.check_in.asc())
    res = await db.execute(query)
    reservations = res.scalars().all()

    rows = []
    for idx, r in enumerate(reservations):
        gobj = None
        if r.guest_id:
            gres = await db.execute(select(Guest).where(Guest.id == r.guest_id))
            gobj = gres.scalar_one_or_none()
        if gobj:
            rows.append({
                "no": idx + 1,
                "full_name": gobj.full_name or "",
                "nationality": gobj.nationality or "",
                "passport_number": gobj.passport_number or "",
                "date_of_issue": gobj.date_of_issue.isoformat() if gobj.date_of_issue else "",
                "arrived_from": gobj.arrived_from or "",
                "visa_number": gobj.visa_number or "",
                "visa_validity": gobj.visa_validity.isoformat() if gobj.visa_validity else "",
                "check_in": r.check_in.isoformat() if r.check_in else "",
                "check_out": r.check_out.isoformat() if r.check_out else "",
            })

    is_pt = (language or "PT").upper() == "PT"
    today_str = date.today().strftime("%d/%m/%Y")
    MONTHS_PT = ["Janeiro","Fevereiro","Março","Abril","Maio","Junho","Julho","Agosto","Setembro","Outubro","Novembro","Dezembro"]
    MONTHS_EN = ["January","February","March","April","May","June","July","August","September","October","November","December"]
    month_name = MONTHS_PT[start_date.month - 1] if is_pt else MONTHS_EN[start_date.month - 1]

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    def add_centered(text, bold=False, size=11):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        return p

    if is_pt:
        add_centered("REPÚBLICA DE MOÇAMBIQUE", bold=True, size=11)
        add_centered("MINISTÉRIO DO INTERIOR", bold=True, size=10)
        add_centered("SERVIÇO NACIONAL DE MIGRAÇÃO", bold=True, size=10)
        add_centered("DIRECÇÃO PROVINCIAL DE MIGRAÇÃO - INHAMBANE", size=10)
        add_centered("DEPARTAMENTO DE OPERAÇÕES MIGRATÓRIAS", size=10)
        add_centered("REPARTIÇÃO DO MOVIMENTO MIGRATÓRIO", size=10)
        add_centered("POSTO DE TRAVESSIA AÉREO DE VILANKULO", size=10)
        p = add_centered("BOLETIM DE ALOJAMENTO - ARTIGO 40 DA LEI N°23/2022 DE 29 DE DEZEMBRO", bold=True, size=9)
    else:
        add_centered("REPUBLIC OF MOZAMBIQUE", bold=True, size=11)
        add_centered("MINISTRY OF INTERIOR", bold=True, size=10)
        add_centered("NATIONAL MIGRATION SERVICE", bold=True, size=10)
        add_centered("PROVINCIAL MIGRATION DIRECTION - INHAMBANE", size=10)
        add_centered("MIGRATION OPERATIONS DEPARTMENT", size=10)
        add_centered("MIGRATORY MOVEMENT SECTION", size=10)
        add_centered("VILANKULO AIR PORT OF ENTRY", size=10)
        p = add_centered("ACCOMMODATION BULLETIN - ARTICLE 40 OF LAW N°23/2022 OF DECEMBER 29", bold=True, size=9)
    p.paragraph_format.space_after = Pt(4)

    # Meta info table
    meta = doc.add_table(rows=2, cols=2)
    meta.style = "Table Grid"
    meta.cell(0, 0).text = f"{'LOCALIZAÇÃO' if is_pt else 'LOCATION'}: {lodge_location or ''}"
    meta.cell(0, 1).text = f"{'ESTABELECIMENTO' if is_pt else 'ESTABLISHMENT'}: {lodge_name or ''}"
    meta.cell(1, 0).text = f"{'MÊS' if is_pt else 'MONTH'}: {month_name}"
    meta.cell(1, 1).text = f"{'DATA' if is_pt else 'DATE'}: {today_str}"
    for row in meta.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

    # Headers
    if is_pt:
        headers = ["N/O","NOME COMPLETO","NACIONALIDADE","Nº PASSAPORTE","DATA EMISSÃO","PROVENIÊNCIA","VISTO Nº","VALIDADE","ENTRADA","SAÍDA"]
    else:
        headers = ["N/O","FULL NAME","NATIONALITY","PASSPORT Nº","DATE OF ISSUE","ARRIVED FROM","VISA Nº","VALIDITY","ENTRY","EXIT"]

    col_widths = [Cm(1.0), Cm(4.5), Cm(3.0), Cm(2.5), Cm(2.2), Cm(3.0), Cm(2.2), Cm(2.2), Cm(2.0), Cm(2.0)]

    table = doc.add_table(rows=1, cols=10)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = table.rows[0].cells
    for i, (cell, hdr) in enumerate(zip(hdr_row, headers)):
        cell.width = col_widths[i]
        cell.text = hdr
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(8)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row_data in rows:
        values = [
            str(row_data["no"]), row_data["full_name"], row_data["nationality"],
            row_data["passport_number"], row_data["date_of_issue"], row_data["arrived_from"],
            row_data["visa_number"], row_data["visa_validity"], row_data["check_in"], row_data["check_out"],
        ]
        row_cells = table.add_row().cells
        for i, (cell, val) in enumerate(zip(row_cells, values)):
            cell.width = col_widths[i]
            cell.text = val
            for para in cell.paragraphs:
                if i == 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(8)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    fname = f"boletim_{(lodge_name or 'lodge').replace(' ','_')}_{month_name.lower()}_{start_date.year}.docx"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'}
    )

@router.post("/reservations")
async def create_reservation(data: ReservationCreate, db: AsyncSession = Depends(get_db)):
    # Auto calculate total_usd if 0
    total = data.total_usd
    if total == 0 and data.rate_per_night_usd > 0 and data.check_out > data.check_in:
        nights = (data.check_out - data.check_in).days
        total = float(data.rate_per_night_usd) * nights
    
    dump = data.model_dump()
    dump["total_usd"] = total
    r = Reservation(**dump)
    db.add(r)
    await db.commit()
    await db.refresh(r)
    return {"id": str(r.id), "status": "created"}

@router.put("/reservations/{res_id}")
async def update_reservation(res_id: uuid.UUID, data: ReservationCreate, db: AsyncSession = Depends(get_db)):
    res = await db.execute(select(Reservation).where(Reservation.id == res_id))
    r = res.scalar_one_or_none()
    if not r:
        raise HTTPException(status_code=404, detail="Reservation not found")
    for k, v in data.model_dump().items():
        setattr(r, k, v)
    await db.commit()
    return {"id": str(r.id), "status": "updated"}

@router.delete("/reservations/{res_id}", status_code=204)
async def delete_reservation(res_id: uuid.UUID, db: AsyncSession = Depends(get_db)):
    res = await db.execute(select(Reservation).where(Reservation.id == res_id))
    r = res.scalar_one_or_none()
    if r:
        await db.delete(r)
        await db.commit()
    return None


# -----------------------------------------------------------------------------
# Staff Task Board CRUD
# -----------------------------------------------------------------------------
@router.get("/tasks", response_model=List[LodgeTaskOut])
async def list_tasks(
    area: Optional[str] = None,
    complete: Optional[bool] = None,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    query = select(LodgeTask)
    if lodge_id:
        query = query.where(LodgeTask.lodge_id == lodge_id)
        
    query = query.order_by(LodgeTask.due_date.asc())
    if area and area != "all":
        query = query.where(LodgeTask.area == area)
    if complete is not None:
        query = query.where(LodgeTask.is_complete == complete)
    res = await db.execute(query)
    return res.scalars().all()

@router.post("/tasks", response_model=LodgeTaskOut, status_code=201)
async def create_task(
    data: LodgeTaskCreate,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    t = LodgeTask(**data.model_dump())
    if lodge_id:
        t.lodge_id = lodge_id
    db.add(t)
    await db.commit()
    await db.refresh(t)
    return t

@router.put("/tasks/{task_id}", response_model=LodgeTaskOut)
async def update_task(
    task_id: uuid.UUID,
    data: LodgeTaskCreate,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    stmt = select(LodgeTask).where(LodgeTask.id == task_id)
    if lodge_id:
        stmt = stmt.where(LodgeTask.lodge_id == lodge_id)
    res = await db.execute(stmt)
    t = res.scalar_one_or_none()
    if not t:
        raise HTTPException(status_code=404, detail="Task not found")
    
    was_complete = t.is_complete
    for k, v in data.model_dump().items():
        setattr(t, k, v)
    
    # If newly marked complete and has recurrence, auto-generate next task
    if not was_complete and t.is_complete and t.recurrence != "none" and t.due_date:
        next_date = t.due_date
        if t.recurrence == "daily":
            next_date += timedelta(days=1)
        elif t.recurrence == "weekly":
            next_date += timedelta(days=7)
        elif t.recurrence == "monthly":
            next_date += timedelta(days=30)
        
        next_task = LodgeTask(
            title=t.title,
            assigned_to=t.assigned_to,
            area=t.area,
            due_date=next_date,
            recurrence=t.recurrence,
            is_complete=False,
            notes=t.notes,
            lodge_id=lodge_id
        )
        db.add(next_task)
    
    await db.commit()
    await db.refresh(t)
    return t

@router.delete("/tasks/{task_id}", status_code=204)
async def delete_task(
    task_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    stmt = select(LodgeTask).where(LodgeTask.id == task_id)
    if lodge_id:
        stmt = stmt.where(LodgeTask.lodge_id == lodge_id)
    res = await db.execute(stmt)
    t = res.scalar_one_or_none()
    if t:
        await db.delete(t)
        await db.commit()
    return None


# -----------------------------------------------------------------------------
# Incidents CRUD
# -----------------------------------------------------------------------------
@router.get("/incidents", response_model=List[IncidentOut])
async def list_incidents(
    resolved: Optional[bool] = None,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    query = select(Incident)
    if lodge_id:
        query = query.where(Incident.lodge_id == lodge_id)
        
    query = query.order_by(
        desc(Incident.severity == "critical"),
        desc(Incident.created_at)
    )
    if resolved is not None:
        query = query.where(Incident.is_resolved == resolved)
    res = await db.execute(query)
    return res.scalars().all()

@router.post("/incidents", response_model=IncidentOut, status_code=201)
async def create_incident(
    data: IncidentCreate,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    i = Incident(**data.model_dump())
    if lodge_id:
        i.lodge_id = lodge_id
    db.add(i)
    await db.commit()
    await db.refresh(i)
    return i

@router.put("/incidents/{inc_id}/resolve", response_model=IncidentOut)
async def resolve_incident(
    inc_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    stmt = select(Incident).where(Incident.id == inc_id)
    if lodge_id:
        stmt = stmt.where(Incident.lodge_id == lodge_id)
    res = await db.execute(stmt)
    i = res.scalar_one_or_none()
    if not i:
        raise HTTPException(status_code=404, detail="Incident not found")
    i.is_resolved = True
    i.resolved_at = datetime.now()
    await db.commit()
    await db.refresh(i)
    return i

@router.delete("/incidents/{inc_id}", status_code=204)
async def delete_incident(
    inc_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    stmt = select(Incident).where(Incident.id == inc_id)
    if lodge_id:
        stmt = stmt.where(Incident.lodge_id == lodge_id)
    res = await db.execute(stmt)
    i = res.scalar_one_or_none()
    if i:
        await db.delete(i)
        await db.commit()
    return None


# -----------------------------------------------------------------------------
# Daily Log CRUD
# -----------------------------------------------------------------------------
@router.get("/daily-log", response_model=List[DailyLogOut])
async def list_daily_logs(
    limit: int = 30,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    query = select(DailyLog)
    if lodge_id:
        query = query.where(DailyLog.lodge_id == lodge_id)
    res = await db.execute(query.order_by(desc(DailyLog.log_date)).limit(limit))
    return res.scalars().all()

@router.post("/daily-log", response_model=DailyLogOut, status_code=201)
async def create_or_update_daily_log(
    data: DailyLogCreate,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    stmt = select(DailyLog).where(DailyLog.log_date == data.log_date)
    if lodge_id:
        stmt = stmt.where(DailyLog.lodge_id == lodge_id)
    res = await db.execute(stmt)
    existing = res.scalar_one_or_none()
    if existing:
        for k, v in data.model_dump().items():
            setattr(existing, k, v)
        await db.commit()
        await db.refresh(existing)
        return existing
    else:
        dl = DailyLog(**data.model_dump())
        if lodge_id:
            dl.lodge_id = lodge_id
        db.add(dl)
        await db.commit()
        await db.refresh(dl)
        return dl
